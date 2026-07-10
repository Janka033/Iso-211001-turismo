"""Verificación lado a lado del formato de los 7 documentos de ``samples/``
+ verificación determinística de los CÓDIGOS del sistema documental por tenant.

Fase 1 (offline, ``samples/``): todos los documentos pertenecen al MISMO
sistema documental (formato Felipe): fuente Calibri 11, encabezado con
Código/Revisión/paginación "Página N de M", bloque de firmas
Elaborado/Revisado/Aprobado y tabla CONTROL DE CAMBIOS (docx); encabezado con
código en la hoja (xlsx, tipo MT: sin firmas por diseño).

Fase 2 (DB, requiere backend/.env): para cada tenant con documentos activos,
resuelve el código como lo hace generation (``onboarding_data.document_codes``,
sin entrada => [PENDIENTE]) y FALLA si (a) dos documentos activos del mismo
tenant resuelven al MISMO código confirmado, o (b) un código confirmado no
está en el catálogo confirmado de Felipe (FELIPE_REAL.md). Los [PENDIENTE]
son estado válido (código aún no confirmado), no fallo.

Uso:  python scripts/check_format_consistency.py [--skip-db]
Sale con código 1 si algo rompe la consistencia.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
SAMPLES = ROOT / "samples"
sys.path.insert(0, str(ROOT / "backend"))

# Catálogo CONFIRMADO del SGS real de Felipe (fuente: FELIPE_REAL.md).
# politica_seguridad / plan_emergencias / gestion_incidentes NO tienen código
# confirmado: deben resolver a [PENDIENTE], nunca a un valor adivinado.
FELIPE_CONFIRMED_CODES: dict[str, str] = {
    "matriz_riesgos": "MT-04",
    "manual_perfiles_cargos": "MA-02",
    "manual_inspeccion_equipos": "MA-03",
    "comunicacion_participacion_consulta": "PR-07",
}

# document_type -> (código esperado, portada esperada per taxonomía)
DOCX_EXPECTED: dict[str, tuple[str, bool]] = {
    "politica_seguridad": ("PO-01", False),  # PO: solo encabezado, sin portada
    "plan_emergencias": ("PL-01", True),
    "gestion_incidentes": ("PR-02", True),
    "manual_perfiles_cargos": ("MA-02", True),
    "comunicacion_participacion_consulta": ("PR-07", True),
    "manual_inspeccion_equipos": ("MA-03", True),
}


def _header_text(doc: Document) -> str:
    parts = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            parts.append(p.text)
        for t in section.header.tables:
            for row in t.rows:
                parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def _body_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def check_docx(name: str, code: str, cover: bool) -> list[str]:
    doc = Document(str(SAMPLES / f"{name}.docx"))
    problems: list[str] = []

    font = doc.styles["Normal"].font
    if font.name != "Calibri":
        problems.append(f"fuente Normal = {font.name!r}, se esperaba Calibri")
    if font.size is None or font.size.pt != 11:
        problems.append(f"tamaño Normal = {font.size}, se esperaba 11pt")

    header = _header_text(doc)
    if f"Código: {code}" not in header:
        problems.append(f"encabezado sin 'Código: {code}'")
    if "Revisión: 01" not in header:
        problems.append("encabezado sin 'Revisión: 01'")
    # El campo PAGE debe llevar resultado cacheado: sin él, los visores
    # muestran "Página  de 1" (bug corregido en FelipeDocxBuilder).
    if "Página 1 de" not in header:
        problems.append("encabezado sin paginación 'Página N de M' (campo PAGE vacío)")

    body = _body_text(doc)
    for required in ("Firmas", "Elaborado por", "Revisado por", "Aprobado por"):
        if required not in body:
            problems.append(f"sin bloque de firmas ('{required}')")
            break
    if "CONTROL DE CAMBIOS" not in body:
        problems.append("sin tabla CONTROL DE CAMBIOS")

    # La portada imprime el código centrado en el cuerpo; sin portada, no.
    has_cover_code = f"Código: {code}" in "\n".join(p.text for p in doc.paragraphs)
    if cover and not has_cover_code:
        problems.append("se esperaba portada (código centrado) y no está")
    if not cover and has_cover_code:
        problems.append("tiene portada y el tipo PO va sin portada")

    return problems


def check_xlsx(name: str, code: str) -> list[str]:
    wb = load_workbook(str(SAMPLES / f"{name}.xlsx"))
    ws = wb["Matriz de Riesgos"]
    problems: list[str] = []
    header_cells = [
        str(ws.cell(row=r, column=1).value or "") for r in range(1, 7)
    ]
    if not any(f"Código: {code}" in c for c in header_cells):
        problems.append(f"hoja sin 'Código: {code}' en el bloque de identidad")
    if not any("Revisión: 01" in c for c in header_cells):
        problems.append("hoja sin 'Revisión: 01'")
    return problems


PENDING_CODE = "[PENDIENTE: código por confirmar]"


def check_tenant_codes() -> int:
    """Fase 2: códigos por tenant contra el catálogo confirmado (via DB).

    Replica la resolución de generation: ``document_codes.get(document_type)``
    o [PENDIENTE]. Falla si hay códigos confirmados duplicados en el mismo
    tenant o un código fuera del catálogo confirmado de Felipe.
    """
    from app.core.supabase import get_service_client  # requiere backend/.env

    client = get_service_client()
    docs = (
        client.table("documents").select("tenant_id, document_type").execute()
    ).data or []
    onboardings = (
        client.table("onboarding_data").select("tenant_id, data").execute()
    ).data or []
    codes_by_tenant = {
        row["tenant_id"]: (row.get("data") or {}).get("document_codes") or {}
        for row in onboardings
    }

    active: dict[str, set[str]] = {}
    for d in docs:
        active.setdefault(d["tenant_id"], set()).add(d["document_type"])

    failures = 0
    for tenant_id, doc_types in sorted(active.items()):
        codes = codes_by_tenant.get(tenant_id, {})
        resolved = {dt: (codes.get(dt) or "").strip() or PENDING_CODE
                    for dt in sorted(doc_types)}

        seen: dict[str, str] = {}
        for dt, code in resolved.items():
            flag = ""
            if code != PENDING_CODE:
                if code in seen:
                    flag = f"FALLA: código duplicado con {seen[code]}"
                    failures += 1
                elif FELIPE_CONFIRMED_CODES.get(dt) != code:
                    flag = "FALLA: código fuera del catálogo confirmado"
                    failures += 1
                else:
                    seen[code] = dt
            print(f"  [{'FALLA' if flag else 'OK'}] {tenant_id[:8]}… {dt}: {code}"
                  + (f"  <- {flag}" if flag else ""))
    return failures


def main() -> int:
    failures = 0
    print("Fase 1 — formato de samples/:")
    for name, (code, cover) in DOCX_EXPECTED.items():
        problems = check_docx(name, code, cover)
        status = "OK" if not problems else "FALLA"
        print(f"  [{status}] {name} ({code})")
        for p in problems:
            print(f"         - {p}")
        failures += bool(problems)

    problems = check_xlsx("matriz_riesgos", "MT-04")
    status = "OK" if not problems else "FALLA"
    print(f"  [{status}] matriz_riesgos (MT-04, xlsx)")
    for p in problems:
        print(f"         - {p}")
    failures += bool(problems)

    if "--skip-db" not in sys.argv:
        print("\nFase 2 — códigos por tenant contra el catálogo confirmado:")
        failures += check_tenant_codes()

    print()
    if failures:
        print(f"{failures} inconsistencia(s) en el sistema documental.")
        return 1
    print("Sistema documental consistente (formato + códigos).")
    return 0


if __name__ == "__main__":
    sys.exit(main())
