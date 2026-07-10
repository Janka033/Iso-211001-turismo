"""Generador .docx del Procedimiento de gestión de incidentes (PR-02, numeral 8.3).

Dos partes:
  - Procedimiento: contenido inteligente (variables de la IA, validadas).
  - Formato de reporte de incidentes: estructura FIJA en blanco (anexo), para
    diligenciar a mano. Sus etiquetas no las pone la IA.

Tipo PR de la taxonomía: portada + contenido + firmas + CONTROL DE CAMBIOS,
vía ``FelipeDocxBuilder``. Código provisional PR-02 (PR-01 es el documento que
define el sistema de Felipe): confirmar numeración real con él.
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydantic import BaseModel

from app.modules.generation.generators.base import (
    DocumentGenerator,
    ResolvedField,
    resolve_code,
)
from app.modules.generation.generators.felipe_docx import FelipeDocxBuilder
from app.modules.generation.schemas import IncidentManagementVariables

# Campos del FORMATO (fijos, en blanco para diligenciar).
_FORM_FIELDS: list[str] = [
    "N.° de reporte",
    "Fecha",
    "Hora",
    "Lugar",
    "Actividad",
    "Tipo de evento (casi-accidente / incidente / accidente)",
    "Personas involucradas",
    "Descripción del evento",
    "Lesiones o daños",
    "Acciones inmediatas tomadas",
    "Causa probable",
    "Acciones correctivas propuestas",
    "Reportado por",
    "Firma / cargo",
]


class IncidentManagementGenerator(DocumentGenerator):
    template_version = "gestion-incidentes-docx-v2"
    engine = "docx"

    def _render(
        self,
        resolved: dict[str, ResolvedField],
        variables: BaseModel,
        document_code: str | None,
    ) -> bytes:
        assert isinstance(variables, IncidentManagementVariables)

        def val(key: str) -> str:
            f = resolved[key]
            return f.value if isinstance(f.value, str) else "\n".join(f.value)

        def items(key: str) -> list[str]:
            f = resolved[key]
            return f.value if isinstance(f.value, list) else [f.value]

        approval_date = (variables.approval_date or "").strip() or (
            "[PENDIENTE: fecha de aprobación]"
        )
        legal_rep = (variables.legal_representative or "").strip()

        b = FelipeDocxBuilder(
            code=resolve_code(document_code),
            title="Procedimiento de gestión de incidentes",
            company=val("company_name"),
            norm_reference="NTC-ISO 21101 — numeral 8.3",
            approval_date=approval_date,
        )
        doc = b.doc

        b.section("1. Objetivo", val("objective"))
        b.section("2. Alcance", val("scope"))
        b.bullet_list("3. Clasificación de eventos", items("incident_classification"))

        b.heading("4. Procedimiento")
        for i, step in enumerate(items("procedure_steps"), start=1):
            doc.add_paragraph(f"{i}. {step}")

        b.section("5. Responsable", val("responsible"))
        b.section("6. Acciones correctivas", val("corrective_actions_process"))
        b.section("7. Conservación de registros", val("record_retention"))
        b.section("8. Revisión del procedimiento", val("review_frequency"))

        b.signatures(approved=(legal_rep, "Representante legal" if legal_rep else ""))
        b.change_control()

        # --- ANEXO: FORMATO DE REPORTE (estructura fija, en blanco) ------
        doc.add_page_break()
        form_title = doc.add_heading("FORMATO DE REPORTE DE INCIDENTES", level=1)
        form_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for label in _FORM_FIELDS:
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(label).bold = True
            cells[1].text = ""  # en blanco, para diligenciar

        return b.render()
