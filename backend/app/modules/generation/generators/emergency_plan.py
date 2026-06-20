"""Generador .docx del Plan de respuesta a emergencias (numeral 8.2).

Estructura fija. La IA aporta los escenarios y los campos planos; las celdas
sin dato quedan ``[PENDIENTE: ...]``, nunca inventadas.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pydantic import BaseModel

from app.modules.generation.generators.base import (
    DocumentGenerator,
    ResolvedField,
    resolve_text,
)
from app.modules.generation.schemas import EmergencyPlanVariables, EmergencyScenario

# Columnas de la tabla de escenarios: (clave, encabezado).
_SCENARIO_COLUMNS: list[tuple[str, str]] = [
    ("scenario", "Escenario de emergencia"),
    ("activity", "Actividad"),
    ("response_steps", "Pasos de respuesta"),
    ("responsible", "Responsable"),
    ("resources", "Recursos"),
]


class EmergencyPlanGenerator(DocumentGenerator):
    template_version = "plan-emergencias-docx-v1"
    engine = "docx"
    custom_fields = frozenset({"emergency_scenarios"})

    def _extra_pending(
        self, variables: BaseModel, required_fields: set[str] | None
    ) -> list[str]:
        assert isinstance(variables, EmergencyPlanVariables)
        if variables.emergency_scenarios or not self._is_required(
            "emergency_scenarios", required_fields
        ):
            return []
        return ["emergency_scenarios"]

    def _render(
        self, resolved: dict[str, ResolvedField], variables: BaseModel
    ) -> bytes:
        assert isinstance(variables, EmergencyPlanVariables)

        def val(key: str) -> str:
            f = resolved[key]
            return f.value if isinstance(f.value, str) else "\n".join(f.value)

        def items(key: str) -> list[str]:
            f = resolved[key]
            return f.value if isinstance(f.value, list) else [f.value]

        doc = Document()
        title = doc.add_heading("PLAN DE RESPUESTA A EMERGENCIAS", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(val("company_name"))
        run.bold = True
        run.font.size = Pt(13)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run("Conforme a la NTC-ISO 21101 — numeral 8.2").italic = True

        self._section(doc, "1. Objetivo", val("general_objective"))
        self._section(doc, "2. Alcance", val("scope"))

        # --- Escenarios de emergencia (tabla) --------------------------
        doc.add_heading("3. Escenarios de emergencia y respuesta", level=1)
        self._scenario_table(doc, variables.emergency_scenarios)

        self._section(
            doc, "4. Protocolo de comunicación", val("communication_protocol")
        )

        doc.add_heading("5. Contactos de emergencia", level=1)
        for contact in items("emergency_contacts"):
            doc.add_paragraph(contact, style="List Bullet")

        self._section(doc, "6. Rutas y recursos de evacuación", val("evacuation_resources"))
        self._section(doc, "7. Coordinación con servicios externos", val("external_coordination"))
        self._section(doc, "8. Capacitación y simulacros", val("training_drills"))
        self._section(doc, "9. Revisión y actualización", val("review_frequency"))

        doc.add_heading("10. Aprobación", level=1)
        doc.add_paragraph(f"Representante legal: {val('legal_representative')}")
        doc.add_paragraph(f"Fecha de aprobación: {val('approval_date')}")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _scenario_table(doc: Document, scenarios: list[EmergencyScenario]) -> None:
        rows = scenarios or [EmergencyScenario()]  # fila vacía => [PENDIENTE]
        field_desc = {
            k: v.description or k for k, v in EmergencyScenario.model_fields.items()
        }
        table = doc.add_table(rows=1, cols=len(_SCENARIO_COLUMNS))
        table.style = "Light Grid Accent 1"
        for col, (_, title) in enumerate(_SCENARIO_COLUMNS):
            table.rows[0].cells[col].paragraphs[0].add_run(title).bold = True
        for entry in rows:
            cells = table.add_row().cells
            for col, (key, _) in enumerate(_SCENARIO_COLUMNS):
                text, _ = resolve_text(getattr(entry, key), field_desc[key])
                cells[col].text = text

    @staticmethod
    def _section(doc: Document, heading: str, body: str) -> None:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
