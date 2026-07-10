"""Layout .docx del sistema documental real de Felipe (TOURS AND ADVENTURES).

Un solo lugar para el FORMATO compartido por la familia Manual/Procedimiento/
Plan/Programa/…: Calibri 11, encabezado con código/empresa/revisión/paginación,
portada, secciones, bloque de firmas y tabla CONTROL DE CAMBIOS. Los generadores
por documento son delgados: mapean sus Variables a secciones y delegan aquí el
layout. No inventa contenido — recibe strings ya resueltos ([PENDIENTE] incluido).
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

CALIBRI = "Calibri"


def apply_base_style(doc: Document) -> None:
    """Estilo base del sistema documental: Normal en Calibri 11.

    python-docx deja ``Normal.font.name = None`` (ambiguo: algunos visores lo
    muestran Arial). Fijarlo explícitamente da consistencia. Reutilizable por los
    generadores existentes para el fix Arial→Calibri.
    """
    normal = doc.styles["Normal"].font
    normal.name = CALIBRI
    normal.size = Pt(11)
    # Los headings heredan del TEMA de Word (Calibri Light); el sistema
    # documental de Felipe es Calibri en todo el texto, títulos incluidos.
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].font.name = CALIBRI


def _page_field(paragraph, instruction: str) -> None:
    """Inserta un campo de Word (PAGE / NUMPAGES) en el párrafo.

    El campo lleva las CUATRO partes (begin / instrText / separate / resultado
    cacheado / end). Sin ``separate`` + resultado, varios visores (y Word hasta
    recalcular campos) muestran el campo VACÍO — era el bug "Página de 1" del
    encabezado. El "1" cacheado se actualiza al repaginar/imprimir.
    """
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, cached, end):
        run._r.append(element)


class FelipeDocxBuilder:
    """Construye un .docx con el formato estándar de Felipe."""

    def __init__(
        self,
        *,
        code: str,
        title: str,
        company: str,
        norm_reference: str,
        approval_date: str,
        revision: str = "01",
        cover: bool = True,
    ) -> None:
        """``cover=False`` para los tipos "solo encabezado" de la taxonomía
        (PO Política, PT, FT, MT, FR, AC): llevan encabezado, firmas y control
        de cambios, pero SIN portada."""
        self.doc = Document()
        apply_base_style(self.doc)
        self.code = code
        self.title = title
        self.company = company
        self.norm_reference = norm_reference
        self.approval_date = approval_date
        self.revision = revision
        self._narrow_margins()
        self._build_header()
        if cover:
            self._build_cover()

    # -- Montaje inicial --------------------------------------------------

    def _narrow_margins(self) -> None:
        for section in self.doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

    def _build_header(self) -> None:
        header = self.doc.sections[0].header
        table = header.add_table(rows=1, cols=3, width=Inches(7.1))
        table.style = "Table Grid"
        left, center, right = table.rows[0].cells
        left.paragraphs[0].add_run("LOGO").bold = True

        cp = center.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(f"{self.title}\n").bold = True
        cp.add_run(f"Código: {self.code}")

        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.add_run(f"{self.company}\n")
        rp.add_run(f"Aprobación: {self.approval_date}\n")
        rp.add_run(f"Revisión: {self.revision}")

        pager = header.add_paragraph()
        pager.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pager.add_run("Página ")
        _page_field(pager, "PAGE")
        pager.add_run(" de ")
        _page_field(pager, "NUMPAGES")

    def _build_cover(self) -> None:
        doc = self.doc
        for _ in range(4):
            doc.add_paragraph()
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run(self.title.upper())
        run.bold = True
        run.font.size = Pt(22)

        code_p = doc.add_paragraph()
        code_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        code_p.add_run(f"Código: {self.code}").font.size = Pt(12)

        doc.add_paragraph()
        comp = doc.add_paragraph()
        comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = comp.add_run(self.company)
        crun.bold = True
        crun.font.size = Pt(15)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(
            f"Fecha de aprobación: {self.approval_date}\nRevisión: {self.revision}"
        )

        norm = doc.add_paragraph()
        norm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        norm.add_run(f"Conforme a la {self.norm_reference}").italic = True

        doc.add_page_break()

    # -- Cuerpo (lo usa cada generador) -----------------------------------

    def section(self, heading: str, body: str, level: int = 1) -> None:
        self.doc.add_heading(heading, level=level)
        self.doc.add_paragraph(body)

    def bullet_list(self, heading: str, items: list[str], level: int = 1) -> None:
        self.doc.add_heading(heading, level=level)
        for item in items:
            self.doc.add_paragraph(item, style="List Bullet")

    def heading(self, text: str, level: int = 1) -> None:
        self.doc.add_heading(text, level=level)

    def table(self, headers: list[str], rows: list[list[str]]) -> None:
        """Tabla con encabezado en negrita. ``rows`` ya viene resuelto a texto."""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for col, header in enumerate(headers):
            table.rows[0].cells[col].paragraphs[0].add_run(header).bold = True
        for row in rows:
            cells = table.add_row().cells
            for col, value in enumerate(row):
                cells[col].text = value

    # -- Cierre estándar --------------------------------------------------

    def signatures(
        self,
        elaborated: tuple[str, str] = ("", ""),
        reviewed: tuple[str, str] = ("", ""),
        approved: tuple[str, str] = ("", ""),
    ) -> None:
        """Bloque de firmas Elaborado/Revisado/Aprobado (nombre, cargo, fecha).

        Cada par es (nombre, cargo); vacío => [PENDIENTE] (no se inventan firmantes).
        """
        self.doc.add_heading("Firmas", level=1)
        headers = ["", "Nombre", "Cargo", "Fecha"]
        table = self.doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for col, header in enumerate(headers):
            table.rows[0].cells[col].paragraphs[0].add_run(header).bold = True
        for label, (name, role) in (
            ("Elaborado por", elaborated),
            ("Revisado por", reviewed),
            ("Aprobado por", approved),
        ):
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(label).bold = True
            cells[1].text = name.strip() or "[PENDIENTE: nombre del firmante]"
            cells[2].text = role.strip() or "[PENDIENTE: cargo del firmante]"
            cells[3].text = "[PENDIENTE: fecha de firma]"

    def change_control(self, entries: list[tuple[str, str, str]] | None = None) -> None:
        """Tabla CONTROL DE CAMBIOS (descripción, fecha, revisión)."""
        self.doc.add_heading("CONTROL DE CAMBIOS", level=1)
        rows = entries or [("Emisión inicial", self.approval_date, self.revision)]
        table = self.doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for col, header in enumerate(["Descripción del cambio", "Fecha", "Revisión"]):
            table.rows[0].cells[col].paragraphs[0].add_run(header).bold = True
        for description, date, revision in rows:
            cells = table.add_row().cells
            cells[0].text = description
            cells[1].text = date
            cells[2].text = revision

    def render(self) -> bytes:
        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()
