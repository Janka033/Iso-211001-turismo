"""Generador .docx del Alcance del SGSTA (numeral 4.3).

Documento corto (1 página) que declara qué cubre el sistema de gestión. El
molde del kit MinCIT es una tabla con el título y la declaración del alcance;
aquí se conserva esa esencia con el layout del sistema documental de Felipe
(tipo corto: encabezado, firmas y control de cambios, SIN portada).

Es la variable clave del árbol: los demás documentos citan este alcance.
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pydantic import BaseModel

from app.modules.generation.generators.base import (
    DocumentGenerator,
    ResolvedField,
    resolve_code,
)
from app.modules.generation.generators.felipe_docx import FelipeDocxBuilder
from app.modules.generation.schemas import AlcanceVariables


class AlcanceGenerator(DocumentGenerator):
    template_version = "alcance-sgsta-docx-v1"
    engine = "docx"

    def _render(
        self,
        resolved: dict[str, ResolvedField],
        variables: BaseModel,
        document_code: str | None,
    ) -> bytes:
        assert isinstance(variables, AlcanceVariables)

        def val(key: str) -> str:
            field = resolved[key]
            return field.value if isinstance(field.value, str) else "\n".join(field.value)

        def items(key: str) -> list[str]:
            field = resolved[key]
            return field.value if isinstance(field.value, list) else [field.value]

        # Datos de aprobación: reales si el cliente los dio; si no, [PENDIENTE]
        # con el mismo marcador que usan los demás documentos.
        approval_date = (variables.approval_date or "").strip() or (
            "[PENDIENTE: fecha de aprobación]"
        )
        legal_rep = (variables.legal_representative or "").strip()

        b = FelipeDocxBuilder(
            code=resolve_code(document_code),
            title="Alcance del SGSTA",
            company=val("company_name"),
            norm_reference="NTC-ISO 21101 — numeral 4.3",
            approval_date=approval_date,
            cover=False,  # documento corto: solo encabezado, sin portada
        )
        doc = b.doc

        title = doc.add_heading(
            "ALCANCE DEL SISTEMA DE GESTIÓN DE SEGURIDAD DE TURISMO DE AVENTURA — SGSTA",
            level=0,
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(val("company_name"))
        run.bold = True
        run.font.size = Pt(13)

        nit_p = doc.add_paragraph()
        nit_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nit_p.add_run(f"NIT: {val('nit')}")

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run("Conforme a la NTC-ISO 21101 — numeral 4.3").italic = True

        # La celda central del molde MinCIT: la declaración es la sección 1 y
        # el resto la desglosa (servicios, ubicación, entorno, aplicabilidad).
        b.section("1. Declaración del alcance", val("scope_statement"))
        b.bullet_list("2. Servicios y actividades cubiertas", items("activities"))
        b.section("3. Ubicación de la operación", val("locations_detail"))
        b.section("4. Características del entorno de operación", val("site_characteristics"))
        b.section("5. Aplicabilidad de los requisitos de la norma", val("norm_exclusions"))

        b.signatures(approved=(legal_rep, "Representante legal" if legal_rep else ""))
        b.change_control()
        return b.render()
