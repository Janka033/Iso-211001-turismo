"""Generador .docx del Acta de compromiso de la alta dirección (numeral 5.1).

Calca el molde del kit MinCIT: objeto, asistentes, la tabla fija de
requisitos 5.1 → compromisos de la dirección, y la constancia de firma. La
tabla es contenido normativo FIJO (los ejemplos del molde son los exigidos);
la IA solo aporta identidad, firmante, ciudad y fecha.
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydantic import BaseModel

from app.modules.generation.generators.base import (
    DocumentGenerator,
    ResolvedField,
    resolve_code,
)
from app.modules.generation.generators.felipe_docx import FelipeDocxBuilder
from app.modules.generation.schemas import ActaCompromisoVariables

# Tabla del molde MinCIT (5.1): requisito de la norma → compromiso de la alta
# dirección. Contenido normativo fijo: no lo escribe la IA.
_COMMITMENTS: tuple[tuple[str, str], ...] = (
    (
        "Asegurar que se establezcan la política de seguridad de turismo de "
        "aventura y los objetivos de seguridad y que sean compatibles con la "
        "dirección estratégica de la organización.",
        "Establecer misión, visión, políticas y objetivos de seguridad.",
    ),
    (
        "Asegurarse de la integración de los requisitos del sistema de gestión "
        "de la seguridad de turismo de aventura en los procesos empresariales "
        "de la organización.",
        "Comprensión adecuada del contexto, comprensión de las necesidades y "
        "expectativas de las partes interesadas, mejora continua.",
    ),
    (
        "Asegurar la disponibilidad de los recursos necesarios para el sistema "
        "de gestión de la seguridad de turismo de aventura.",
        "Presupuesto.\nCaracterizaciones de procesos.",
    ),
    (
        "Comunicar la importancia de una gestión eficaz de la seguridad y de "
        "ajustarse a los requisitos del sistema de gestión de la seguridad de "
        "turismo de aventura.",
        "Toma de conciencia.\nPrograma de capacitación y formación.\nMatriz de "
        "comunicación.",
    ),
    (
        "Asegurar que el sistema de gestión de la seguridad de turismo de "
        "aventura logre los resultados previstos.",
        "Mediante la revisión del SGSTA y verificación del cumplimiento de los "
        "objetivos del sistema de gestión de seguridad establecidos por la "
        "organización.",
    ),
    (
        "Dirigir y apoyar a las personas para que contribuyan a la eficacia "
        "del sistema de gestión de la seguridad de turismo de aventura.",
        "Estableciendo roles, responsabilidades y autoridades en el sistema de "
        "gestión de la seguridad de turismo de aventura.",
    ),
    (
        "Promover la mejora continua.",
        "Estableciendo el compromiso con las auditorías internas.\nRealizando "
        "revisión por la dirección según la frecuencia establecida.\nTomando "
        "las acciones correctivas y de mejora pertinentes, de manera oportuna.",
    ),
    (
        "Apoyar a otras funciones de gestión pertinentes para demostrar su "
        "liderazgo en sus esferas de responsabilidad.",
        "Involucrando a los líderes de proceso en la revisión por la dirección "
        "periódica.\nEstableciendo roles, responsabilidades y autoridades en "
        "el sistema de gestión de seguridad.",
    ),
)


class ActaCompromisoGenerator(DocumentGenerator):
    template_version = "acta-compromiso-docx-v1"
    engine = "docx"
    static_sections = (
        "Tabla de compromisos de la alta dirección: los 8 requisitos del "
        "numeral 5.1 (política y objetivos, integración en los procesos, "
        "recursos, comunicación, resultados, apoyo a las personas, mejora "
        "continua, liderazgo en otras funciones) con el compromiso del kit "
        "MinCIT para cada uno",
    )

    def _render(
        self,
        resolved: dict[str, ResolvedField],
        variables: BaseModel,
        document_code: str | None,
    ) -> bytes:
        assert isinstance(variables, ActaCompromisoVariables)

        def val(key: str) -> str:
            field = resolved[key]
            return field.value if isinstance(field.value, str) else "\n".join(field.value)

        acta_date = (variables.acta_date or "").strip() or (
            "[PENDIENTE: fecha de firma del acta]"
        )
        signer = (variables.signer_name or "").strip()

        b = FelipeDocxBuilder(
            code=resolve_code(document_code),
            title="Acta de compromiso",
            company=val("company_name"),
            norm_reference="NTC-ISO 21101 — numeral 5.1",
            approval_date=acta_date,
            cover=False,  # documento corto: solo encabezado, sin portada
        )
        doc = b.doc

        title = doc.add_heading("ACTA DE COMPROMISO DE LA ALTA DIRECCIÓN", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        place = doc.add_paragraph()
        place.alignment = WD_ALIGN_PARAGRAPH.CENTER
        place.add_run(f"{val('city')}, {acta_date}")

        b.table(
            ["OBJETO", "ASISTENTES"],
            [
                [
                    "Demostrar el compromiso y liderazgo de la alta dirección "
                    "con el sistema de gestión de seguridad de turismo de "
                    f"aventura de {val('company_name')}.",
                    f"Nombre: {val('signer_name')}\n"
                    f"Teléfono: {val('signer_phone')}\n"
                    f"Correo electrónico: {val('signer_email')}\n"
                    f"Cargo: {val('signer_role')}",
                ]
            ],
        )

        doc.add_paragraph(
            "Conforme a lo establecido por la NTC-ISO 21101:2020 — Requisitos "
            "turismo de aventura. Sistemas de gestión de seguridad, la "
            f"organización {val('company_name')}, a través de la alta dirección "
            "y la presente acta, manifiesta su compromiso con el sistema de "
            "gestión de seguridad de turismo de aventura, promoviendo el "
            "cumplimiento de los siguientes ítems:"
        )

        b.table(
            ["REQUISITOS", "COMPROMISO ALTA DIRECCIÓN"],
            [[req, com] for req, com in _COMMITMENTS],
        )

        doc.add_paragraph(
            f"Se firma como constancia de aceptación el {acta_date}."
        )

        signer_role = (variables.signer_role or "").strip() if signer else ""
        b.signatures(approved=(signer, signer_role))
        b.change_control()
        return b.render()
