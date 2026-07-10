"""Generador .docx de la Política de seguridad (PO-01, numeral 5.2).

La ESTRUCTURA del documento es fija (este código). La IA solo aporta los
valores de las variables, ya validados por Pydantic y resueltos a texto o
``[PENDIENTE: ...]`` por la clase base.

Tipo PO de la taxonomía: "solo encabezado + contenido libre" — lleva el
encabezado, firmas y CONTROL DE CAMBIOS del sistema documental, SIN portada.
Código provisional PO-01: confirmar numeración real con Felipe.
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pydantic import BaseModel

from app.modules.generation.generators.base import DocumentGenerator, ResolvedField
from app.modules.generation.generators.felipe_docx import FelipeDocxBuilder
from app.modules.generation.schemas import SecurityPolicyVariables


class SecurityPolicyGenerator(DocumentGenerator):
    template_version = "politica-seguridad-docx-v3"
    engine = "docx"

    def _render(
        self, resolved: dict[str, ResolvedField], variables: BaseModel
    ) -> bytes:
        assert isinstance(variables, SecurityPolicyVariables)

        def val(key: str) -> str:
            field = resolved[key]
            return field.value if isinstance(field.value, str) else "\n".join(field.value)

        def items(key: str) -> list[str]:
            field = resolved[key]
            return field.value if isinstance(field.value, list) else [field.value]

        # Datos de aprobación: reales si el cliente los dio; si no, [PENDIENTE]
        # con el MISMO marcador que usan MA-02/PR-07/MA-03.
        approval_date = (variables.approval_date or "").strip() or (
            "[PENDIENTE: fecha de aprobación]"
        )
        legal_rep = (variables.legal_representative or "").strip()

        b = FelipeDocxBuilder(
            code="PO-01",
            title="Política de seguridad",
            company=val("company_name"),
            norm_reference="NTC-ISO 21101 — numeral 5.2",
            approval_date=approval_date,
            cover=False,  # tipo PO: solo encabezado, sin portada
        )
        doc = b.doc

        # --- Identidad (contenido libre del tipo PO) ---------------------
        title = doc.add_heading("POLÍTICA DE SEGURIDAD", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(val("company_name"))
        run.bold = True
        run.font.size = Pt(13)

        nit_p = doc.add_paragraph()
        nit_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nit_p.add_run(f"NIT: {val('nit')}")

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run("Conforme a la NTC-ISO 21101 — numeral 5.2").italic = True

        # --- Cuerpo -----------------------------------------------------
        # Las secciones 4-7 son los 4 principios del 5.2 en su orden: (a)
        # apropiación del propósito, (b) marco para objetivos, (c) cumplimiento
        # de requisitos, (d) mejora continua. La 3 (compromiso) los introduce.
        b.section("1. Alcance", val("scope"))
        b.bullet_list("2. Actividades de turismo de aventura", items("activities"))
        b.section("3. Compromiso de la alta dirección", val("management_commitment"))
        b.section(
            "4. Apropiación del propósito organizacional", val("purpose_alignment")
        )
        b.bullet_list("5. Objetivos de seguridad", items("safety_objectives"))
        b.section(
            "6. Compromiso de cumplimiento legal y reglamentario",
            val("legal_commitment"),
        )
        b.section("7. Compromiso de mejora continua", val("continuous_improvement"))
        b.section("8. Comunicación y disponibilidad", val("communication"))

        # --- Cierre estándar del sistema documental ----------------------
        b.signatures(approved=(legal_rep, "Representante legal" if legal_rep else ""))
        b.change_control()
        return b.render()
