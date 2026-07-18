"""Tests del Procedimiento de control de la información documentada (7.5.1).

La metodología y los tiempos de retención son fijos (molde MinCIT); la IA
solo aporta identidad, cargo responsable, almacenamiento y aprobación.
"""

import io

from docx import Document

from app.modules.generation.generators.control_documental import (
    ControlDocumentalGenerator,
)
from app.modules.generation.generators.factory import get_spec
from app.modules.generation.schemas import ControlDocumentalVariables


def _text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_control_documental_full_document():
    variables = ControlDocumentalVariables(
        company_name="Rafting Fonce SAS",
        responsible_role="Gerente",
        storage_description="Electrónicos en la nube; físicos en carpetas en la oficina",
        legal_representative="Ana Ruiz",
        approval_date="2026-01-15",
    )
    result = ControlDocumentalGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    # Metodología fija del molde.
    assert "control de la información documentada" in txt
    assert "Documentos obsoletos" in txt
    assert "dos (2)" in txt  # retención SGSTA
    assert "veinte (20)" in txt  # retención SST
    assert "Todos los colaboradores del PST" in txt
    # Personalización del cliente.
    assert "responsabilidad de: Gerente" in txt
    assert "Electrónicos en la nube" in txt
    assert "Ana Ruiz" in txt
    assert "CONTROL DE CAMBIOS" in txt


def test_control_documental_empty_marks_pending():
    result = ControlDocumentalGenerator().generate(ControlDocumentalVariables())
    assert "company_name" in result.pending_fields
    assert "responsible_role" in result.pending_fields
    assert "storage_description" in result.pending_fields
    txt = _text(result.content)
    assert "[PENDIENTE:" in txt
    # La metodología fija sale aunque falte todo lo demás.
    assert "veinte (20)" in txt


def test_control_documental_responsible_fills_steps_table():
    variables = ControlDocumentalVariables(
        company_name="X SAS",
        responsible_role="Coordinadora SGSTA",
        storage_description="Nube y archivador",
    )
    txt = _text(ControlDocumentalGenerator().generate(variables).content)
    # Sección 5 + pasos personalizables (2, 3, 6, 8, 9).
    assert txt.count("Coordinadora SGSTA") >= 6


def test_control_documental_registered_in_factory():
    spec = get_spec("control_informacion_documentada")
    assert spec is not None
    assert spec.numeral == "7.5.1"
    assert spec.variables_model is ControlDocumentalVariables
    assert spec.rag_numerales == ("7.5.1", "7.5.2", "7.5.3")
    assert spec.generator.static_sections
