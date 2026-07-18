"""Tests del Acta de compromiso (5.1): generador 3/8 de la ruta guiada.

La tabla requisitos→compromisos es contenido normativo fijo de la plantilla;
la IA solo aporta identidad, firmante, ciudad y fecha.
"""

import io

from docx import Document

from app.modules.generation.generators.acta_compromiso import ActaCompromisoGenerator
from app.modules.generation.generators.factory import get_spec
from app.modules.generation.schemas import ActaCompromisoVariables


def _text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_acta_full_document():
    variables = ActaCompromisoVariables(
        company_name="Rafting Fonce SAS",
        city="San Gil",
        acta_date="15/01/2026",
        signer_name="Ana Ruiz",
        signer_role="Gerente",
        signer_phone="300 111 2233",
        signer_email="ana@raftingfonce.co",
    )
    result = ActaCompromisoGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    assert "ACTA DE COMPROMISO DE LA ALTA DIRECCIÓN" in txt
    assert "San Gil, 15/01/2026" in txt
    assert "Ana Ruiz" in txt
    assert "ana@raftingfonce.co" in txt
    # La tabla normativa fija del molde MinCIT está completa (8 requisitos).
    assert "Promover la mejora continua." in txt
    assert "Establecer misión, visión, políticas y objetivos de seguridad." in txt
    assert "Se firma como constancia de aceptación el 15/01/2026." in txt
    assert "CONTROL DE CAMBIOS" in txt


def test_acta_empty_marks_pending():
    result = ActaCompromisoGenerator().generate(ActaCompromisoVariables())
    assert "company_name" in result.pending_fields
    assert "signer_name" in result.pending_fields
    txt = _text(result.content)
    assert "[PENDIENTE:" in txt
    # El contenido normativo fijo sale aunque falte todo lo demás.
    assert "Promover la mejora continua." in txt


def test_acta_optional_contact_does_not_pend():
    """Con la checklist real (0049), teléfono/correo/fecha son opcionales."""
    required = {"company_name", "city", "signer_name", "signer_role"}
    variables = ActaCompromisoVariables(
        company_name="X SAS",
        city="San Gil",
        signer_name="Ana Ruiz",
        signer_role="Gerente",
    )
    result = ActaCompromisoGenerator().generate(variables, required_fields=required)
    assert result.pending_fields == []
    # La fecha de firma sin dato queda visible como pendiente en el acta.
    assert "[PENDIENTE: fecha de firma del acta]" in _text(result.content)


def test_acta_registered_in_factory():
    spec = get_spec("acta_compromiso")
    assert spec is not None
    assert spec.numeral == "5.1"
    assert spec.variables_model is ActaCompromisoVariables
    assert spec.rag_numerales == ("5.1",)
    assert spec.generator.static_sections  # el evaluador conoce la tabla fija
