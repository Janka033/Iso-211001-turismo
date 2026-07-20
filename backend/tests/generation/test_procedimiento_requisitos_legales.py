"""Tests del Procedimiento de requisitos legales (6.1.3).

Toda la metodología (objetivo, alcance, las 4 definiciones y las 7 actividades)
es fija del molde MinCIT; la IA solo aporta identidad, cargo responsable y
aprobación. Complementa la Matriz de requisitos legales (MT-03).
"""

import io

from docx import Document

from app.modules.generation.generators.factory import get_spec
from app.modules.generation.generators.procedimiento_requisitos_legales import (
    RequisitosLegalesProcedimientoGenerator,
)
from app.modules.generation.schemas import ProcedimientoRequisitosLegalesVariables


def _text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_requisitos_legales_full_document():
    variables = ProcedimientoRequisitosLegalesVariables(
        company_name="Rafting Fonce SAS",
        responsible_role="Gerente",
        legal_representative="Ana Ruiz",
        approval_date="2026-01-15",
    )
    result = RequisitosLegalesProcedimientoGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    # Metodología fija del molde.
    assert "Disposición legislativa provisional" in txt  # definición Decreto Ley
    assert "ESCNNA" in txt  # generalidades (actividad 1)
    assert "cada seis meses" in txt  # frecuencia de actualización (actividad 4)
    assert "plan de intervención" in txt  # actividades 6 y 7
    assert "Matriz de requisitos legales" in txt  # registro
    # Personalización del cliente.
    assert "responsabilidad de: Gerente" in txt
    assert "Ana Ruiz" in txt
    assert "CONTROL DE CAMBIOS" in txt


def test_requisitos_legales_empty_marks_pending():
    result = RequisitosLegalesProcedimientoGenerator().generate(
        ProcedimientoRequisitosLegalesVariables()
    )
    assert "company_name" in result.pending_fields
    assert "responsible_role" in result.pending_fields
    txt = _text(result.content)
    assert "[PENDIENTE:" in txt
    # La metodología fija sale aunque falte todo lo demás.
    assert "Disposición legislativa provisional" in txt


def test_requisitos_legales_responsible_fills_all_activities():
    variables = ProcedimientoRequisitosLegalesVariables(
        company_name="X SAS",
        responsible_role="Coordinadora SGSTA",
    )
    txt = _text(RequisitosLegalesProcedimientoGenerator().generate(variables).content)
    # Sección 4 (responsable) + las 7 actividades: el cargo se repite >= 8 veces.
    assert txt.count("Coordinadora SGSTA") >= 8


def test_requisitos_legales_registered_in_factory():
    spec = get_spec("procedimiento_requisitos_legales")
    assert spec is not None
    assert spec.numeral == "6.1.3"
    assert spec.variables_model is ProcedimientoRequisitosLegalesVariables
    assert spec.generator.static_sections
