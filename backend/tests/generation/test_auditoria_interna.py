"""Tests del Procedimiento de auditoría interna (9.2).

Toda la metodología (objetivo, alcance, 20 definiciones GTC-ISO 19011,
aclaraciones —perfil del auditor, frecuencia anual— y las 11 actividades) es
fija del molde MinCIT; la IA solo aporta identidad, cargo responsable y
aprobación.
"""

import io

from docx import Document

from app.modules.generation.generators.auditoria_interna import (
    AuditoriaInternaGenerator,
)
from app.modules.generation.generators.factory import get_spec
from app.modules.generation.schemas import AuditoriaInternaVariables


def _text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_auditoria_interna_full_document():
    variables = AuditoriaInternaVariables(
        company_name="Rafting Fonce SAS",
        responsible_role="Gerente",
        legal_representative="Ana Ruiz",
        approval_date="2026-01-15",
    )
    result = AuditoriaInternaGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    # Metodología fija del molde.
    assert "GTC-ISO 19011" in txt
    assert "No conformidad" in txt  # una de las 20 definiciones
    assert "una (1) auditoría interna en el año" in txt  # frecuencia mínima
    assert "no puede auditar su propio trabajo" in txt  # independencia
    assert "Auditor interno" in txt  # responsable fijo de la ejecución
    assert "Informe de auditoría" in txt  # registro de la actividad 10
    # Personalización del cliente.
    assert "responsabilidad de: Gerente" in txt
    assert "Ana Ruiz" in txt
    assert "CONTROL DE CAMBIOS" in txt


def test_auditoria_interna_empty_marks_pending():
    result = AuditoriaInternaGenerator().generate(AuditoriaInternaVariables())
    assert "company_name" in result.pending_fields
    assert "responsible_role" in result.pending_fields
    txt = _text(result.content)
    assert "[PENDIENTE:" in txt
    # La metodología fija sale aunque falte todo lo demás.
    assert "GTC-ISO 19011" in txt


def test_auditoria_interna_responsible_fills_programming_steps():
    variables = AuditoriaInternaVariables(
        company_name="X SAS",
        responsible_role="Coordinadora SGSTA",
    )
    txt = _text(AuditoriaInternaGenerator().generate(variables).content)
    # Sección 5 (responsable) + las 4 actividades de programación (1-4).
    assert txt.count("Coordinadora SGSTA") >= 5


def test_auditoria_interna_registered_in_factory():
    spec = get_spec("auditoria_interna")
    assert spec is not None
    assert spec.numeral == "9.2"
    assert spec.variables_model is AuditoriaInternaVariables
    assert spec.generator.static_sections
