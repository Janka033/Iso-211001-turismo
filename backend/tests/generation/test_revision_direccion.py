"""Tests del Procedimiento de revisión por la dirección (9.3).

Toda la metodología (objetivo, alcance, definiciones, entradas y las 6
actividades) es fija del molde MinCIT; la IA solo aporta identidad, cargo
responsable y aprobación.
"""

import io

from docx import Document

from app.modules.generation.generators.factory import get_spec
from app.modules.generation.generators.revision_direccion import (
    RevisionDireccionGenerator,
)
from app.modules.generation.schemas import RevisionDireccionVariables


def _text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_revision_direccion_full_document():
    variables = RevisionDireccionVariables(
        company_name="Rafting Fonce SAS",
        responsible_role="Gerente",
        legal_representative="Ana Ruiz",
        approval_date="2026-01-15",
    )
    result = RevisionDireccionGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    # Metodología fija del molde.
    assert "Idoneidad" in txt  # una de las definiciones
    assert "una vez al año" in txt  # frecuencia mínima de la revisión
    assert "resultados de la auditoría" in txt.lower()  # entrada obligatoria
    assert "Alta dirección" in txt  # responsable fijo
    assert "adecuación, idoneidad y eficacia" in txt  # conclusiones
    # Personalización del cliente.
    assert "responsabilidad de: Gerente" in txt
    assert "Ana Ruiz" in txt
    assert "CONTROL DE CAMBIOS" in txt


def test_revision_direccion_empty_marks_pending():
    result = RevisionDireccionGenerator().generate(RevisionDireccionVariables())
    assert "company_name" in result.pending_fields
    assert "responsible_role" in result.pending_fields
    txt = _text(result.content)
    assert "[PENDIENTE:" in txt
    assert "Mejora continua" in txt  # la metodología fija sale igual


def test_revision_direccion_registered_in_factory():
    spec = get_spec("revision_direccion")
    assert spec is not None
    assert spec.numeral == "9.3"
    assert spec.variables_model is RevisionDireccionVariables
    assert spec.rag_numerales == ("9.3", "9.1", "9.2")
    assert spec.generator.static_sections
