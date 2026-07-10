"""Tests de los generadores .docx de plan de emergencias (8.2) y gestión de
incidentes (8.3)."""

import io

from docx import Document

from app.modules.generation.generators.emergency_plan import EmergencyPlanGenerator
from app.modules.generation.generators.incident_management import (
    IncidentManagementGenerator,
)
from app.modules.generation.schemas import (
    EmergencyPlanVariables,
    EmergencyScenario,
    IncidentManagementVariables,
)


def _doc(content: bytes) -> Document:
    return Document(io.BytesIO(content))


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


# --- Plan de emergencias (8.2) ----------------------------------------------


def test_emergency_plan_full_no_pending():
    variables = EmergencyPlanVariables(
        company_name="Aventura Andina SAS",
        scope="Rafting y canopy en Santander.",
        general_objective="Responder eficazmente a emergencias.",
        emergency_scenarios=[
            EmergencyScenario(
                scenario="Volcamiento de balsa",
                activity="Rafting",
                response_steps="Activar rescate acuático y dar primeros auxilios.",
                responsible="Guía líder",
                resources="Cuerdas de rescate, botiquín",
            )
        ],
        communication_protocol="Cadena: guía -> coordinador -> 123.",
        emergency_contacts=["Bomberos 119", "Ambulancia 125"],
        evacuation_resources="Ruta fluvial hasta el punto de embarque.",
        external_coordination="Hospital regional y Defensa Civil.",
        training_drills="Simulacro semestral.",
        review_frequency="Anual.",
        legal_representative="María Gómez",
        approval_date="2026-06-18",
    )
    result = EmergencyPlanGenerator().generate(variables, document_code='PL-01')
    assert result.pending_fields == []
    doc = _doc(result.content)
    text = _all_text(doc)
    assert "PLAN DE RESPUESTA A EMERGENCIAS" in text
    assert "Volcamiento de balsa" in text
    # Cierre estándar del sistema documental (formato Felipe).
    assert "Firmas" in text
    assert "CONTROL DE CAMBIOS" in text
    # El CONTENIDO (párrafos) no lleva pendientes. La fecha de FIRMA queda
    # [PENDIENTE] por diseño (se firma en papel) y vive en la tabla de firmas.
    paragraphs = "\n".join(p.text for p in doc.paragraphs)
    assert "[PENDIENTE" not in paragraphs


def test_emergency_plan_empty_scenarios_pending():
    variables = EmergencyPlanVariables(company_name="Sin Datos SAS")
    result = EmergencyPlanGenerator().generate(variables)
    assert "emergency_scenarios" in result.pending_fields
    assert "general_objective" in result.pending_fields
    text = _all_text(_doc(result.content))
    assert "[PENDIENTE:" in text


# --- Gestión de incidentes (8.3) --------------------------------------------


def test_incident_full_includes_fixed_form():
    variables = IncidentManagementVariables(
        company_name="Aventura Andina SAS",
        scope="Todas las actividades.",
        objective="Gestionar incidentes y prevenir su recurrencia.",
        incident_classification=["Casi-accidente", "Incidente", "Accidente"],
        procedure_steps=["Reportar", "Registrar", "Investigar", "Acciones correctivas"],
        responsible="Coordinador HSE",
        corrective_actions_process="Análisis de causa raíz y verificación.",
        record_retention="5 años en archivo digital.",
        review_frequency="Anual.",
        legal_representative="María Gómez",
        approval_date="2026-06-18",
    )
    result = IncidentManagementGenerator().generate(variables)
    assert result.pending_fields == []
    text = _all_text(_doc(result.content))
    assert "PROCEDIMIENTO DE GESTIÓN DE INCIDENTES" in text
    # El procedimiento numera los pasos.
    assert "1. Reportar" in text
    # El FORMATO fijo está presente con sus etiquetas.
    assert "FORMATO DE REPORTE DE INCIDENTES" in text
    assert "N.° de reporte" in text
    assert "Causa probable" in text


def test_incident_missing_fields_pending_but_form_still_present():
    variables = IncidentManagementVariables(company_name="X SAS")
    result = IncidentManagementGenerator().generate(variables)
    assert "objective" in result.pending_fields
    assert "procedure_steps" in result.pending_fields
    text = _all_text(_doc(result.content))
    # El formato (estructura fija) no depende de la IA: siempre aparece.
    assert "FORMATO DE REPORTE DE INCIDENTES" in text
    assert "[PENDIENTE:" in text
