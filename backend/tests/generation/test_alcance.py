"""Tests del Alcance del SGSTA (4.3): primer generador nuevo de la ruta guiada.

Mismo patrón que los demás generadores: documento completo, pendientes
visibles cuando falta dato, y registro en la factory con su RAG multi-numeral.
"""

import io

from docx import Document

from app.modules.generation.generators.alcance import AlcanceGenerator
from app.modules.generation.generators.factory import get_spec
from app.modules.generation.schemas import AlcanceVariables


def _text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_alcance_full_document():
    variables = AlcanceVariables(
        company_name="Rafting Fonce SAS",
        nit="901.234.567-8",
        activities=["Rafting", "Torrentismo"],
        scope_statement=(
            "El sistema de gestión de seguridad para turismo de aventura de "
            "Rafting Fonce SAS cubre la prestación de los servicios de rafting "
            "y torrentismo en el río Fonce, San Gil, Santander."
        ),
        locations_detail="Santander, San Gil — río Fonce y cascada Juan Curí.",
        site_characteristics="Zona rural, ribera de río, área de influencia de reserva.",
        norm_exclusions="Aplican todos los requisitos de la norma.",
        legal_representative="Ana Ruiz",
        approval_date="2026-01-15",
    )
    result = AlcanceGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    assert "ALCANCE DEL SISTEMA DE GESTIÓN" in txt
    assert "Rafting Fonce SAS" in txt
    assert "río Fonce" in txt
    assert "Torrentismo" in txt
    assert "Aplican todos los requisitos" in txt
    assert "CONTROL DE CAMBIOS" in txt
    assert "Firmas" in txt
    assert "Ana Ruiz" in txt
    assert "2026-01-15" in txt
    assert Document(io.BytesIO(result.content)).styles["Normal"].font.name == "Calibri"


def test_alcance_empty_marks_pending():
    result = AlcanceGenerator().generate(AlcanceVariables())
    assert "scope_statement" in result.pending_fields
    assert "activities" in result.pending_fields
    assert "norm_exclusions" in result.pending_fields
    assert "[PENDIENTE:" in _text(result.content)


def test_alcance_optional_fields_do_not_pend():
    """Con la checklist real (0047), firma y fecha son opcionales: sin dato
    salen neutros, no como pendiente."""
    required = {
        "company_name",
        "nit",
        "activities",
        "scope_statement",
        "locations_detail",
        "site_characteristics",
        "norm_exclusions",
    }
    variables = AlcanceVariables(
        company_name="X SAS",
        nit="900.000.000-1",
        activities=["Rafting"],
        scope_statement="Cubre rafting en el río Fonce.",
        locations_detail="Santander, San Gil.",
        site_characteristics="Zona rural.",
        norm_exclusions="Aplican todos.",
    )
    result = AlcanceGenerator().generate(variables, required_fields=required)
    assert result.pending_fields == []
    txt = _text(result.content)
    # La fecha de aprobación sin dato sí queda visible como pendiente (control
    # de cambios), igual que en los demás documentos del sistema.
    assert "[PENDIENTE: fecha de aprobación]" in txt


def test_alcance_registered_in_factory():
    spec = get_spec("alcance_sgsta")
    assert spec is not None
    assert spec.numeral == "4.3"
    assert spec.variables_model is AlcanceVariables
    # El RAG amplía con el contexto de la organización (4.1) y partes
    # interesadas (4.2), que delimitan qué debe cubrir el alcance.
    assert spec.rag_numerales == ("4.3", "4.1", "4.2")
