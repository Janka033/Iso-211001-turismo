"""Tests del Procedimiento de gestión de riesgos y oportunidades (6.1.1).

La metodología es fija (molde MinCIT); la IA solo aporta identidad, cargo
líder, frecuencia y aprobación.
"""

import io

from docx import Document

from app.modules.generation.generators.factory import get_spec
from app.modules.generation.generators.procedimiento_riesgos import (
    ProcedimientoRiesgosGenerator,
)
from app.modules.generation.schemas import ProcedimientoRiesgosVariables


def _text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_procedimiento_full_document():
    variables = ProcedimientoRiesgosVariables(
        company_name="Rafting Fonce SAS",
        responsible_role="Gerente",
        follow_up_frequency="Semestral",
        legal_representative="Ana Ruiz",
        approval_date="2026-01-15",
    )
    result = ProcedimientoRiesgosGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    # Metodología fija del molde.
    assert "Establecer el paso a paso para la identificación" in txt
    assert "Gestión del riesgo" in txt
    assert "riesgos ambientales" in txt  # nota de riesgos climáticos
    assert "Seguimiento y evaluación" in txt
    # Personalización del cliente.
    assert "liderada por: Gerente" in txt
    assert "frecuencia: Semestral" in txt
    assert "Todos los líderes de procesos" in txt  # responsable fijo del molde
    assert "Ana Ruiz" in txt
    assert "CONTROL DE CAMBIOS" in txt


def test_procedimiento_empty_marks_pending():
    result = ProcedimientoRiesgosGenerator().generate(ProcedimientoRiesgosVariables())
    assert "company_name" in result.pending_fields
    assert "responsible_role" in result.pending_fields
    assert "follow_up_frequency" in result.pending_fields
    txt = _text(result.content)
    assert "[PENDIENTE:" in txt
    # La metodología fija sale aunque falte todo lo demás.
    assert "Seguimiento y evaluación" in txt


def test_procedimiento_responsible_fills_steps_table():
    """El cargo del cliente aparece como responsable de los pasos que el
    molde deja personalizables (no los de 'Todos los líderes de procesos')."""
    variables = ProcedimientoRiesgosVariables(
        company_name="X SAS",
        responsible_role="Coordinador de seguridad",
        follow_up_frequency="Anual",
    )
    txt = _text(ProcedimientoRiesgosGenerator().generate(variables).content)
    assert txt.count("Coordinador de seguridad") >= 5  # sección 5 + pasos 1,2,5,6,7


def test_procedimiento_registered_in_factory():
    spec = get_spec("procedimiento_riesgos_oportunidades")
    assert spec is not None
    assert spec.numeral == "6.1.1"
    assert spec.variables_model is ProcedimientoRiesgosVariables
    assert spec.rag_numerales == ("6.1.1", "A")
    assert spec.generator.static_sections
