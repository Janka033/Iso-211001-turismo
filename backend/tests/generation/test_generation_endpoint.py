"""Tests de integración del endpoint de generación.

Se mockean los bordes externos (Supabase vía repository, IA vía AIProvider);
se ejercita el flujo real del service: RAG → prompt → validación Pydantic →
generador → snapshot.
"""

import pytest

from app.modules.generation import repository, service
from app.modules.inventory import service as inventory_service

TENANT_A = "11111111-1111-1111-1111-111111111111"

PROMPT_CONTENT = (
    "Extrae las variables.\n"
    "Empresa: <<COMPANY_CONTEXT>>\n"
    "Norma: <<NORM_CONTEXT>>\n"
    "Checklist: <<CHECKLIST>>\n"
    "Patrones: <<PATTERNS>>\n"
    "Schema: <<JSON_SCHEMA>>"
)

CHECKLIST = [
    {"field_key": "company_name", "description": "Razón social", "required": True},
    {"field_key": "scope", "description": "Alcance", "required": True},
    {"field_key": "activities", "description": "Actividades", "required": True},
    {"field_key": "management_commitment", "description": "Compromiso", "required": True},
    {"field_key": "safety_objectives", "description": "Objetivos", "required": True},
    {"field_key": "legal_commitment", "description": "Legal", "required": True},
    {"field_key": "continuous_improvement", "description": "Mejora", "required": True},
    {"field_key": "nit", "description": "NIT", "required": False},
]

FULL_OUTPUT = {
    "company_name": "Aventura Andina SAS",
    "nit": "900.123.456-7",
    "scope": "Rafting y canopy en Santander.",
    "activities": ["Rafting", "Canopy"],
    "management_commitment": "La dirección se compromete con la seguridad.",
    "safety_objectives": ["Cero accidentes graves"],
    "legal_commitment": "Cumplir la normativa de turismo.",
    "continuous_improvement": "Revisión semestral.",
    "communication": "Publicada en la sede.",
    "legal_representative": "María Gómez",
    "approval_date": "2026-06-18",
}


class FakeProvider:
    def __init__(self, output: dict):
        self._output = output

    async def embed(self, text: str) -> list[float]:
        return [0.0] * 768

    async def generate_json(self, prompt: str) -> dict:
        # Guardamos el prompt para inspección si hiciera falta.
        self.last_prompt = prompt
        return self._output


@pytest.fixture
def wire(monkeypatch):
    """Cablea el repository y el provider con dobles. Devuelve un dict de captura."""

    captured: dict = {}

    def setup(ai_output: dict):
        provider = FakeProvider(ai_output)
        monkeypatch.setattr(service, "get_ai_provider", lambda: provider)

        monkeypatch.setattr(
            repository, "get_active_prompt",
            lambda dt, token: {"version": "v1", "content": PROMPT_CONTENT},
        )
        monkeypatch.setattr(
            repository, "get_company",
            lambda tid, token: {"name": "Aventura Andina SAS", "nit": "900.123.456-7"},
        )
        monkeypatch.setattr(
            repository, "get_onboarding_data",
            lambda tid, token: {"actividades": ["rafting"]},
        )
        monkeypatch.setattr(repository, "get_checklist", lambda dt, token: CHECKLIST)
        monkeypatch.setattr(repository, "get_generation_patterns", lambda dt, token: [])
        monkeypatch.setattr(
            repository, "match_knowledge_chunks",
            lambda emb, numeral, token, **kw: [
                {"id": "aaaaaaaa-0000-0000-0000-000000000001",
                 "source": "norma", "numeral": "5.2", "content": "La política..."},
            ],
        )
        async def fake_list_items(tid, token, category=None):
            return []

        monkeypatch.setattr(inventory_service, "list_items", fake_list_items)

        monkeypatch.setattr(repository, "next_version", lambda tid, dt, token: 1)
        monkeypatch.setattr(
            repository, "upload_document",
            lambda tid, dt, v, content, token, engine="docx": f"{tid}/{dt}/v{v}.{engine}",
        )

        def fake_insert(record, token):
            captured["record"] = record
            return {"id": "doc-0001", **record}

        monkeypatch.setattr(repository, "insert_document", fake_insert)

        def fake_status(tid, dt, st, comp, token):
            captured["status"] = {"status": st, "completeness": comp}

        monkeypatch.setattr(repository, "upsert_document_status", fake_status)
        return captured

    return setup


def _auth(make_token):
    token = make_token(tenant_id=TENANT_A, role="empresa")
    return {"Authorization": f"Bearer {token}"}


def test_generate_full_document(client, make_token, wire):
    captured = wire(FULL_OUTPUT)
    resp = client.post("/generation/politica_seguridad", headers=_auth(make_token))

    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["document_id"] == "doc-0001"
    assert body["version"] == 1
    assert body["status"] == "generated"
    assert body["completeness"] == 100.0
    assert body["pending_fields"] == []
    assert body["storage_path"].endswith("v1.docx")

    # Snapshot de reproducibilidad completo.
    record = captured["record"]
    assert record["prompt_version"] == "v1"
    assert record["template_version"] == "politica-seguridad-docx-v4"
    assert record["prompt_hash"]
    assert record["rag_chunks_ids"] == ["aaaaaaaa-0000-0000-0000-000000000001"]
    assert record["variables_snapshot"]["company_name"] == "Aventura Andina SAS"


def test_generate_partial_marks_pending(client, make_token, wire):
    wire({"company_name": "Sin Datos SAS"})  # faltan los obligatorios
    resp = client.post("/generation/politica_seguridad", headers=_auth(make_token))

    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert "scope" in body["pending_fields"]
    assert "activities" in body["pending_fields"]
    assert body["completeness"] < 100.0


def test_invalid_ai_json_is_502(client, make_token, wire):
    # activities debe ser lista; un str rompe la validación Pydantic.
    wire({"company_name": "X", "activities": "no-soy-lista"})
    resp = client.post("/generation/politica_seguridad", headers=_auth(make_token))
    assert resp.status_code == 502


def test_unknown_document_type_is_404(client, make_token, wire):
    wire(FULL_OUTPUT)
    resp = client.post("/generation/documento_inexistente", headers=_auth(make_token))
    assert resp.status_code == 404


def test_requires_tenant_claim(client, make_token, wire):
    wire(FULL_OUTPUT)
    token = make_token(tenant_id=None)  # JWT sin tenant_id
    resp = client.post(
        "/generation/politica_seguridad",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert resp.status_code == 401


def test_inventory_items_reach_the_prompt(client, make_token, wire, monkeypatch):
    """Fase 3: el inventario operativo del tenant entra en el contexto del prompt
    para que la IA genere listas dinámicas (N salidas, vehículos, etc.)."""
    captured = wire(FULL_OUTPUT)

    provider = FakeProvider(FULL_OUTPUT)
    monkeypatch.setattr(service, "get_ai_provider", lambda: provider)

    class _Item:
        def __init__(self, category, name, attributes):
            self.category = category
            self.name = name
            self.attributes = attributes

    async def fake_list_items(tid, token, category=None):
        return [
            _Item("salida_emergencia", "Salida Norte", {"punto_encuentro": "Cancha"}),
            _Item("salida_emergencia", "Salida Sur", {"punto_encuentro": "Parqueadero"}),
            _Item("vehiculo", "Cuatrimoto 01", {"placa": "ABC123"}),
        ]

    monkeypatch.setattr(inventory_service, "list_items", fake_list_items)

    resp = client.post("/generation/politica_seguridad", headers=_auth(make_token))
    assert resp.status_code == 201, resp.text

    prompt = provider.last_prompt
    assert "INVENTARIO OPERATIVO REAL" in prompt  # sección prominente
    assert "Salida Norte" in prompt
    assert "Salida Sur" in prompt
    assert "Cuatrimoto 01" in prompt
    assert "ABC123" in prompt
    # Y queda en el snapshot de reproducibilidad indirectamente vía prompt_hash.
    assert captured["record"]["prompt_hash"]


def test_list_document_types(client, make_token):
    resp = client.get("/generation/document-types", headers=_auth(make_token))
    assert resp.status_code == 200
    assert "politica_seguridad" in resp.json()


async def test_rerender_from_snapshot_is_deterministic_and_never_calls_ai(monkeypatch):
    """Re-render: variables del snapshot + plantilla ACTUAL + código del
    catálogo, versión nueva, procedencia conservada y CERO llamadas a la IA.
    Separa extracción (LLM) de render (determinístico)."""
    import io

    from docx import Document

    def _no_ai():
        raise AssertionError("el re-render no debe llamar al provider de IA")

    monkeypatch.setattr(service, "get_ai_provider", _no_ai)
    monkeypatch.setattr(
        repository,
        "get_latest_document",
        lambda tid, dt, token: {
            "id": "doc-old",
            "version": 1,
            "variables_snapshot": FULL_OUTPUT,
            "prompt_hash": "hash-origen",
            "prompt_version": "v9",
            "rag_chunks_ids": ["chunk-1"],
            "model_version": "gemini-x",
        },
    )
    monkeypatch.setattr(
        repository,
        "get_onboarding_data",
        lambda tid, token: {"document_codes": {"politica_seguridad": "PO-99"}},
    )
    monkeypatch.setattr(repository, "get_checklist", lambda dt, token: CHECKLIST)
    monkeypatch.setattr(repository, "next_version", lambda tid, dt, token: 2)

    captured: dict = {}

    def fake_upload(tid, dt, v, content, token, engine="docx"):
        captured["content"] = content
        return f"{tid}/{dt}/v{v}.{engine}"

    monkeypatch.setattr(repository, "upload_document", fake_upload)

    def fake_insert(record, token):
        captured["record"] = record
        return {"id": "doc-new", **record}

    monkeypatch.setattr(repository, "insert_document", fake_insert)
    monkeypatch.setattr(repository, "upsert_document_status", lambda *a, **k: None)

    doc = await service.rerender_from_snapshot("politica_seguridad", TENANT_A, "tok")

    assert doc.version == 2
    assert doc.template_version == "politica-seguridad-docx-v4"
    assert doc.completeness == 100.0
    # Procedencia de las VARIABLES conservada del registro origen.
    record = captured["record"]
    assert record["prompt_hash"] == "hash-origen"
    assert record["prompt_version"] == "v9"
    assert record["rag_chunks_ids"] == ["chunk-1"]
    assert record["model_version"] == "gemini-x"
    # El render actual aplicó el código del catálogo y la paginación cacheada.
    rendered = Document(io.BytesIO(captured["content"]))
    header_parts = [p.text for p in rendered.sections[0].header.paragraphs]
    for table in rendered.sections[0].header.tables:
        for row in table.rows:
            header_parts.extend(c.text for c in row.cells)
    header = "\n".join(header_parts)
    assert "Código: PO-99" in header
    assert "Página 1 de 1" in header


def test_completeness_ignores_checklist_fields_outside_model():
    """La checklist puede ir por delante del modelo de variables (0025: campos
    de levantamiento de PL-01 aún sin variable/plantilla). Esos field_key no
    deben contarse como "llenos" (inflarían el %) ni como pendientes."""
    checklist = CHECKLIST + [
        {"field_key": "brigada_emergencias", "description": "Brigada", "required": True},
    ]
    model_fields = {c["field_key"] for c in CHECKLIST}

    # Todos los obligatorios del modelo llenos => 100, no 7/8.
    assert service._completeness(checklist, [], model_fields) == 100.0
    # Con un obligatorio real pendiente, el % se mide solo sobre el modelo.
    assert service._completeness(checklist, ["scope"], model_fields) == round(6 / 7 * 100, 2)
