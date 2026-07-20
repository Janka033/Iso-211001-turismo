"""Tests del Procedimiento de identificación y selección de personal (5.3).

Toda la metodología (objetivo, alcance, las 11 definiciones, las aclaraciones y
las 7 actividades) es fija del molde MinCIT; la IA solo aporta identidad, cargo
responsable y aprobación. Complementa el Manual de perfiles y cargos (MA-02).
"""

import io

from docx import Document

from app.modules.generation.generators.factory import get_spec
from app.modules.generation.generators.procedimiento_seleccion_personal import (
    SeleccionPersonalGenerator,
)
from app.modules.generation.schemas import ProcedimientoSeleccionPersonalVariables


def _text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_seleccion_personal_full_document():
    variables = ProcedimientoSeleccionPersonalVariables(
        company_name="Rafting Fonce SAS",
        responsible_role="Coordinador de talento humano",
        legal_representative="Ana Ruiz",
        approval_date="2026-01-15",
    )
    result = SeleccionPersonalGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    # Metodología fija del molde.
    assert "Superintendencia Financiera" in txt  # definición de ARL
    assert "No se podrá firmar el contrato" in txt  # aclaración del molde
    assert "reinducción" in txt  # actividad 4
    assert "toma de conciencia" in txt  # actividad 6 (y el objetivo)
    assert "paz y salvo" in txt  # actividad 7 (desvinculación)
    # Personalización del cliente.
    assert "responsabilidad de: Coordinador de talento humano" in txt
    assert "Ana Ruiz" in txt
    assert "CONTROL DE CAMBIOS" in txt


def test_seleccion_personal_empty_marks_pending():
    result = SeleccionPersonalGenerator().generate(
        ProcedimientoSeleccionPersonalVariables()
    )
    assert "company_name" in result.pending_fields
    assert "responsible_role" in result.pending_fields
    txt = _text(result.content)
    assert "[PENDIENTE:" in txt
    # La metodología fija sale aunque falte todo lo demás.
    assert "Superintendencia Financiera" in txt


def test_seleccion_personal_registered_in_factory():
    spec = get_spec("procedimiento_seleccion_personal")
    assert spec is not None
    assert spec.numeral == "5.3"
    assert spec.variables_model is ProcedimientoSeleccionPersonalVariables
    assert spec.generator.static_sections
    # 7.2 (competencia/formación) amplía el contexto normativo del RAG.
    assert "7.2" in spec.rag_numerales
