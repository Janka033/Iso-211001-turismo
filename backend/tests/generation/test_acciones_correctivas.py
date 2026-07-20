"""Tests del Procedimiento de acciones correctivas y de mejora (10.1).

Toda la metodología (objetivo, alcance, definiciones, aclaraciones y las 6
actividades) es fija del molde MinCIT; la IA solo aporta identidad, cargo
responsable y aprobación.
"""

import io

from docx import Document

from app.modules.generation.generators.acciones_correctivas import (
    AccionesCorrectivasGenerator,
)
from app.modules.generation.generators.factory import get_spec
from app.modules.generation.schemas import AccionesCorrectivasVariables


def _text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_acciones_correctivas_full_document():
    variables = AccionesCorrectivasVariables(
        company_name="Rafting Fonce SAS",
        responsible_role="Gerente",
        legal_representative="Ana Ruiz",
        approval_date="2026-01-15",
    )
    result = AccionesCorrectivasGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    # Metodología fija del molde.
    assert "Mejora continua" in txt  # una de las definiciones
    assert "análisis de causas" in txt  # actividad 3
    assert "evaluación de la eficacia" in txt.lower()  # actividad 6
    assert "Cerrada" in txt  # estado de cierre de la acción
    assert "Líder del proceso" in txt  # responsable fijo de análisis/plan
    # Personalización del cliente.
    assert "responsabilidad de: Gerente" in txt
    assert "Ana Ruiz" in txt
    assert "CONTROL DE CAMBIOS" in txt


def test_acciones_correctivas_empty_marks_pending():
    result = AccionesCorrectivasGenerator().generate(AccionesCorrectivasVariables())
    assert "company_name" in result.pending_fields
    assert "responsible_role" in result.pending_fields
    txt = _text(result.content)
    assert "[PENDIENTE:" in txt
    # La metodología fija sale aunque falte todo lo demás.
    assert "No conformidad" in txt


def test_acciones_correctivas_registered_in_factory():
    spec = get_spec("acciones_correctivas_mejora")
    assert spec is not None
    assert spec.numeral == "10.1"
    assert spec.variables_model is AccionesCorrectivasVariables
    assert spec.rag_numerales == ("10.1", "9.2")
    assert spec.generator.static_sections
