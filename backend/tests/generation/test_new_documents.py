"""Tests de los 3 documentos nuevos del sistema documental de Felipe:
MA-02 (perfiles), PR-07 (comunicación) y MA-03 (inspección de equipos).

Mismo patrón que los existentes: generador (estructura + [PENDIENTE]) y un flujo
por el endpoint (motor docx, snapshot, RAG multi-numeral).
"""

import io

from docx import Document

from app.modules.generation import repository, service
from app.modules.generation.generators.communication_procedure import (
    CommunicationProcedureGenerator,
)
from app.modules.generation.generators.equipment_manual import EquipmentManualGenerator
from app.modules.generation.generators.factory import get_spec
from app.modules.generation.generators.profiles_manual import ProfilesManualGenerator
from app.modules.generation.schemas import (
    CommunicationEntry,
    CommunicationProcedureVariables,
    EquipmentInspectionItem,
    EquipmentManualVariables,
    ProfilesManualVariables,
    RoleProfile,
)
from app.modules.inventory import service as inventory_service

TENANT_A = "11111111-1111-1111-1111-111111111111"


def _text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _normal_font(content: bytes) -> str | None:
    return Document(io.BytesIO(content)).styles["Normal"].font.name


# ---------------------------------------------------------------------------
# MA-02 · Manual de perfiles y funciones de cargo
# ---------------------------------------------------------------------------


def test_ma02_full_document():
    variables = ProfilesManualVariables(
        company_name="Aventura Andina SAS",
        objective="Definir perfiles y funciones de cada cargo.",
        scope="Todo el personal de la operación.",
        org_structure="Nivel directivo y nivel operativo.",
        role_profiles=[
            RoleProfile(
                role="Gerente",
                level="Directivo",
                purpose="Dirigir la organización.",
                functions=["Aprobar la política de seguridad"],
                requirements="Profesional con experiencia.",
                reports_to="Junta de socios",
            )
        ],
        legal_representative="María Gómez",
        approval_date="2026-01-15",
    )
    result = ProfilesManualGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    assert "Gerente" in txt
    assert "Aprobar la política de seguridad" in txt
    assert "CONTROL DE CAMBIOS" in txt
    assert "Firmas" in txt
    # "Aprobado por" con el representante legal (consistente con PO-01/PL-01).
    assert "María Gómez" in txt
    assert "2026-01-15" in txt
    assert _normal_font(result.content) == "Calibri"  # formato Felipe


def test_ma02_empty_marks_pending():
    result = ProfilesManualGenerator().generate(ProfilesManualVariables())
    assert "role_profiles" in result.pending_fields
    assert "objective" in result.pending_fields
    assert "[PENDIENTE:" in _text(result.content)


def test_ma02_partial_profile_reports_cell_pendings():
    """Patrón de risk_matrix: los [PENDIENTE] de celda de una fila real se
    reportan uno a uno (visibles para el cliente y el auditor)."""
    variables = ProfilesManualVariables(
        company_name="X SAS",
        objective="Definir perfiles.",
        scope="Todo el personal.",
        org_structure="Directivo y operativo.",
        role_profiles=[RoleProfile(role="Guía")],  # solo el cargo
    )
    result = ProfilesManualGenerator().generate(variables)
    assert "role_profiles" not in result.pending_fields  # hay al menos un perfil
    assert "role_profiles[0].level" in result.pending_fields
    assert "role_profiles[0].functions" in result.pending_fields
    assert "role_profiles[0].role" not in result.pending_fields


def test_null_lists_from_ai_are_tolerated():
    """El prompt permite 'null (o lista vacía)': un null en un campo lista no
    debe tumbar la validación (502), sino tratarse como dato faltante."""
    variables = ProfilesManualVariables.model_validate(
        {"company_name": "X SAS", "objective": None, "role_profiles": None}
    )
    assert variables.role_profiles == []
    assert variables.objective is None

    equip = EquipmentManualVariables.model_validate(
        {"equipment_items": None, "maintenance_types": None}
    )
    assert equip.equipment_items == []
    assert equip.maintenance_types == []


def test_headings_and_title_use_calibri():
    """Todo el sistema documental es Calibri, títulos incluidos (no el
    Calibri Light del tema de Word)."""
    result = ProfilesManualGenerator().generate(ProfilesManualVariables())
    doc = Document(io.BytesIO(result.content))
    assert doc.styles["Title"].font.name == "Calibri"
    assert doc.styles["Heading 1"].font.name == "Calibri"
    assert doc.styles["Heading 2"].font.name == "Calibri"


# ---------------------------------------------------------------------------
# PR-07 · Comunicación, participación y consulta
# ---------------------------------------------------------------------------


def test_pr07_full_document():
    variables = CommunicationProcedureVariables(
        company_name="X SAS",
        objective="Regular la comunicación del SGS.",
        scope="Todo el personal.",
        responsibles="Líder del Sistema de Gestión.",
        communication_matrix=[
            CommunicationEntry(
                topic="Política de seguridad",
                method="Reunión de inducción",
                channel="Presencial",
                audience="Todo el equipo",
                frequency="Anual",
            )
        ],
        participation="Comité paritario y buzón de sugerencias.",
        consultation="Consulta previa a cambios operativos.",
        representation="Representante del personal en el comité.",
        performance_evaluation="Encuesta anual de comunicación.",
        records="Actas de reunión y registros de asistencia.",
        legal_representative="María Gómez",
        approval_date="2026-01-15",
    )
    result = CommunicationProcedureGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    assert "Política de seguridad" in txt
    assert "Reunión de inducción" in txt
    assert "Anual" in txt
    assert "CONTROL DE CAMBIOS" in txt


def test_pr07_empty_marks_pending():
    result = CommunicationProcedureGenerator().generate(
        CommunicationProcedureVariables()
    )
    assert "communication_matrix" in result.pending_fields
    assert "[PENDIENTE:" in _text(result.content)


def test_pr07_partial_row_reports_cell_pendings():
    variables = CommunicationProcedureVariables(
        communication_matrix=[CommunicationEntry(topic="Política de seguridad")],
    )
    result = CommunicationProcedureGenerator().generate(variables)
    assert "communication_matrix" not in result.pending_fields
    assert "communication_matrix[0].frequency" in result.pending_fields
    assert "communication_matrix[0].topic" not in result.pending_fields


# ---------------------------------------------------------------------------
# MA-03 · Inspección y mantenimiento de equipos (usa el equipo REAL por actividad)
# ---------------------------------------------------------------------------


def test_ma03_uses_real_equipment_per_activity():
    variables = EquipmentManualVariables(
        company_name="X SAS",
        objective="Controlar la inspección y el mantenimiento.",
        scope="Todo el equipo de seguridad.",
        role_obligations="Gerencia aprueba; Coordinador programa; Guía inspecciona.",
        equipment_items=[
            EquipmentInspectionItem(
                activity="rafting",
                equipment="Balsa, remos y chalecos",
                state="Ok",
                inspection_type="previa a uso",
                frequency="Cada salida",
            ),
            EquipmentInspectionItem(
                activity="buceo",
                equipment="Regulador, BCD y tanque",
                state="Ok",
                inspection_type="periódica trimestral",
                frequency="Trimestral",
            ),
        ],
        maintenance_types=["Revisión previa a uso", "Revisión periódica trimestral"],
        legal_representative="María Gómez",
        approval_date="2026-01-15",
    )
    result = EquipmentManualGenerator().generate(variables)

    assert result.pending_fields == []
    txt = _text(result.content)
    # El control de equipos refleja el equipo REAL de cada actividad, no genérico.
    assert "rafting" in txt and "Balsa, remos y chalecos" in txt
    assert "buceo" in txt and "Regulador, BCD y tanque" in txt


def test_ma03_empty_marks_pending():
    result = EquipmentManualGenerator().generate(EquipmentManualVariables())
    assert "equipment_items" in result.pending_fields
    assert "[PENDIENTE:" in _text(result.content)


# ---------------------------------------------------------------------------
# RAG multi-numeral declarado en la factory
# ---------------------------------------------------------------------------


def test_new_specs_declare_rag_numerales():
    assert get_spec("manual_perfiles_cargos").rag_numerales == ("5.3", "7.2")
    assert get_spec("comunicacion_participacion_consulta").rag_numerales == ("7.4",)
    assert get_spec("manual_inspeccion_equipos").rag_numerales == ("8.1", "A")


# ---------------------------------------------------------------------------
# Flujo por el endpoint (motor docx) — representativo con MA-02
# ---------------------------------------------------------------------------

MA02_CHECKLIST = [
    {"field_key": "company_name", "description": "Razón social", "required": True},
    {"field_key": "objective", "description": "Objetivo", "required": True},
    {"field_key": "scope", "description": "Alcance", "required": True},
    {"field_key": "org_structure", "description": "Estructura", "required": True},
    {"field_key": "role_profiles", "description": "Perfiles", "required": True},
]

MA02_OUTPUT = {
    "company_name": "Aventura Andina SAS",
    "objective": "Definir perfiles.",
    "scope": "Todo el personal.",
    "org_structure": "Directivo y operativo.",
    "role_profiles": [
        {"role": "Gerente", "level": "Directivo", "functions": ["Aprobar la política"]}
    ],
}


class _Provider:
    async def embed(self, text: str) -> list[float]:
        return [0.0] * 768

    async def generate_json(self, prompt: str) -> dict:
        return MA02_OUTPUT


def test_ma02_endpoint_flow(client, make_token, monkeypatch):
    captured: dict = {}
    monkeypatch.setattr(service, "get_ai_provider", lambda: _Provider())
    monkeypatch.setattr(
        repository, "get_active_prompt",
        lambda dt, token: {"version": "v1", "content": "<<JSON_SCHEMA>>"},
    )
    monkeypatch.setattr(repository, "get_company", lambda tid, token: {"name": "X"})
    monkeypatch.setattr(repository, "get_onboarding_data", lambda tid, token: {})
    monkeypatch.setattr(repository, "get_checklist", lambda dt, token: MA02_CHECKLIST)
    monkeypatch.setattr(repository, "get_generation_patterns", lambda dt, token: [])

    def fake_match(emb, numeral, token, **kw):
        captured.setdefault("numerales", []).append(numeral)
        return []

    monkeypatch.setattr(repository, "match_knowledge_chunks", fake_match)

    async def fake_list_items(tid, token, category=None):
        return []

    monkeypatch.setattr(inventory_service, "list_items", fake_list_items)
    monkeypatch.setattr(repository, "next_version", lambda tid, dt, token: 1)

    def fake_upload(tid, dt, v, content, token, engine="docx"):
        captured["engine"] = engine
        return f"{tid}/{dt}/v{v}.{engine}"

    monkeypatch.setattr(repository, "upload_document", fake_upload)
    monkeypatch.setattr(
        repository, "insert_document", lambda record, token: {"id": "doc-x", **record}
    )
    monkeypatch.setattr(repository, "upsert_document_status", lambda *a, **k: None)

    token = make_token(tenant_id=TENANT_A, role="empresa")
    resp = client.post(
        "/generation/manual_perfiles_cargos",
        headers={"Authorization": f"Bearer {token}"},
    )

    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert captured["engine"] == "docx"
    assert body["storage_path"].endswith("v1.docx")
    assert body["template_version"] == "manual-perfiles-docx-v3"
    assert body["completeness"] == 100.0
    assert captured["numerales"] == ["5.3", "7.2"]  # RAG multi-numeral
