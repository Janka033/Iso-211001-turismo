"""Tests del generador y de la resolución de placeholders (sin servicios externos)."""

import io

from docx import Document

from app.modules.generation.generators.security_policy import SecurityPolicyGenerator
from app.modules.generation.schemas import SecurityPolicyVariables


def _text_of(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs)


def test_full_variables_no_pending():
    variables = SecurityPolicyVariables(
        company_name="Aventura Andina SAS",
        nit="900.123.456-7",
        scope="Operación de rafting y canopy en Santander.",
        activities=["Rafting", "Canopy"],
        management_commitment="La dirección se compromete con la seguridad.",
        purpose_alignment="Alineada con el propósito de brindar turismo seguro.",
        safety_objectives=["Cero accidentes graves", "Capacitación anual"],
        legal_commitment="Cumplir el Decreto 1085 y la normativa de turismo.",
        continuous_improvement="Revisar el sistema cada 6 meses.",
        communication="Publicada en la sede y el sitio web.",
        legal_representative="María Gómez",
        approval_date="2026-06-18",
    )

    result = SecurityPolicyGenerator().generate(variables)

    assert result.pending_fields == []
    text = _text_of(result.content)
    assert "Aventura Andina SAS" in text
    assert "[PENDIENTE" not in text


def test_missing_fields_become_pending_not_invented():
    # Solo razón social; el resto debe quedar como [PENDIENTE], nunca inventado.
    variables = SecurityPolicyVariables(company_name="Sin Datos SAS")

    result = SecurityPolicyGenerator().generate(variables)

    # Campos obligatorios sin dato deben marcarse pendientes.
    assert "scope" in result.pending_fields
    assert "activities" in result.pending_fields
    assert "management_commitment" in result.pending_fields

    text = _text_of(result.content)
    assert "[PENDIENTE: Alcance" in text
    assert "Sin Datos SAS" in text


def test_blank_string_counts_as_pending():
    variables = SecurityPolicyVariables(company_name="   ", scope="\n")
    result = SecurityPolicyGenerator().generate(variables)
    assert "company_name" in result.pending_fields
    assert "scope" in result.pending_fields


def test_optional_missing_is_neutral_not_pendiente():
    # Todos los obligatorios presentes; los opcionales vacíos NO deben gritar
    # [PENDIENTE], sino quedar como "No especificado".
    variables = SecurityPolicyVariables(
        company_name="X SAS",
        scope="s",
        activities=["a"],
        management_commitment="m",
        safety_objectives=["o"],
        legal_commitment="l",
        continuous_improvement="c",
    )
    required = {
        "company_name", "scope", "activities", "management_commitment",
        "safety_objectives", "legal_commitment", "continuous_improvement",
    }
    result = SecurityPolicyGenerator().generate(variables, required)

    assert result.pending_fields == []
    text = _text_of(result.content)
    assert "[PENDIENTE" not in text
    assert "No especificado" in text  # nit / communication / approval_date opcionales


def test_required_missing_is_pendiente_optional_neutral():
    variables = SecurityPolicyVariables(company_name="X SAS")
    required = {"scope", "activities"}  # solo estos obligatorios
    result = SecurityPolicyGenerator().generate(variables, required)

    assert set(result.pending_fields) == {"scope", "activities"}
    text = _text_of(result.content)
    assert "[PENDIENTE: Alcance" in text
    assert "No especificado" in text  # management_commitment, etc. opcionales => neutro


def test_generates_valid_docx_bytes():
    variables = SecurityPolicyVariables(company_name="X SAS")
    result = SecurityPolicyGenerator().generate(variables)
    # Debe abrirse como un .docx válido.
    doc = Document(io.BytesIO(result.content))
    headings = [p.text for p in doc.paragraphs]
    assert any("POLÍTICA DE SEGURIDAD" in h for h in headings)
